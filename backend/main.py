import io
import email
from email import policy
import pypdf
import docx
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from groq import Groq
import os
import json

load_dotenv()

app = FastAPI(title="AIVOA Pharma QA Complaint System")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = Groq(api_key=os.getenv("GROQ_API_KEY"))

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


class ComplaintRequest(BaseModel):
    complaint: str


class EditComplaintRequest(BaseModel):
    current_data: dict
    correction: str


@app.get("/")
def home():
    return {"message": "AIVOA backend is working!"}


def process_complaint_text_with_ai(text: str) -> dict:
    """Core complaint analysis pipeline using Groq Llama 3.3 70B."""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "system",
                "content": """
You are an AI assistant for an enterprise pharmaceutical quality assurance customer complaint management system (API & FDF QA Module).

Extract all relevant details from the customer complaint and assess the risk.

Return ONLY valid JSON.
Do not use Markdown.
Do not add explanations outside the JSON.

Return exactly this JSON structure:

{
    "complaint_source": "",
    "customer_name": "",
    "product_name": "",
    "strength": "",
    "batch_number": "",
    "manufacturing_date": "",
    "expiry_date": "",
    "affected_quantity": "",
    "complaint_type": "",
    "complaint_date": "",
    "complaint_description": "",
    "risk_assessment": {
        "severity": "",
        "priority": "",
        "initial_risk": "",
        "suggested_next_action": ""
    }
}

Guidelines:
- complaint_source: The reporting channel or entity (e.g., Hospital, Clinic, Pharmacy, Wholesaler, Email, Doctor, Patient).
- customer_name: Full name of reporting individual or organization.
- product_name: Commercial or generic drug/API/FDF name.
- strength: Product strength, concentration, or grade (e.g., 500mg, 100 IU/ml, USP Grade).
- batch_number: Batch or Lot number.
- manufacturing_date: Manufacturing date in YYYY-MM-DD or as stated.
- expiry_date: Expiry date in YYYY-MM-DD or as stated.
- affected_quantity: Quantity affected with units (e.g., 120 bottles, 50 kg, 45 vials).
- complaint_type: Classification (e.g., Physical Defect, Packaging Failure, Discoloration, Contamination, Seal Integrity, Adverse Event, Potency Deviation).
- complaint_date: Date complaint was logged/reported (YYYY-MM-DD).
- complaint_description: Detailed summary of the quality issue observed.
- risk_assessment.severity: Must be one of: Low, Medium, High, Critical.
- risk_assessment.priority: Must be one of: P1 - Critical, P2 - High, P3 - Medium, P4 - Low.
- risk_assessment.initial_risk: Concise clinical/QA risk evaluation.
- risk_assessment.suggested_next_action: SOP/CAPA containment action (e.g. quarantine batch, test retain samples, initiate recall investigation).

Do not invent false details. If a field is not mentioned, provide a reasonable empty string or "Not specified".
"""
            },
            {
                "role": "user",
                "content": f"""Customer complaint / document:

{text}"""
            }
        ],
        response_format={"type": "json_object"}
    )

    result = response.choices[0].message.content.strip()
    if result.startswith("```json"):
        result = result[7:]
    if result.startswith("```"):
        result = result[3:]
    if result.endswith("```"):
        result = result[:-3]

    return json.loads(result.strip())


@app.post("/complaint")
def analyze_complaint(request: ComplaintRequest):
    return process_complaint_text_with_ai(request.complaint)


@app.post("/complaint/upload")
async def upload_complaint_document(file: UploadFile = File(...)):
    # 1. Read file bytes
    contents = await file.read()
    file_size = len(contents)

    # 2. Validate max file size (10 MB)
    if file_size > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=400,
            detail=f"File size exceeds 10 MB limit (File size: {file_size / (1024*1024):.2f} MB)."
        )

    if file_size == 0:
        raise HTTPException(
            status_code=400,
            detail="Uploaded file is empty."
        )

    filename = (file.filename or "").lower()
    extracted_text = ""

    # 3. Extract text based on file type
    try:
        if filename.endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(io.BytesIO(contents))
            pages_text = []
            for i, page in enumerate(pdf_reader.pages):
                page_content = page.extract_text()
                if page_content:
                    pages_text.append(page_content)
            extracted_text = "\n\n".join(pages_text).strip()

        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(contents))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            extracted_text = "\n".join(paragraphs).strip()

        elif filename.endswith(".eml"):
            msg = email.message_from_bytes(contents, policy=policy.default)
            subject = msg.get("subject", "")
            from_addr = msg.get("from", "")
            date_hdr = msg.get("date", "")
            body = msg.get_body(preferencelist=('plain', 'html'))
            body_text = body.get_content() if body else ""
            extracted_text = f"Email Subject: {subject}\nFrom: {from_addr}\nDate: {date_hdr}\n\nBody:\n{body_text}".strip()

        else:
            # Fallback for .txt, .json, .csv, and standard text
            try:
                extracted_text = contents.decode("utf-8").strip()
            except UnicodeDecodeError:
                extracted_text = contents.decode("latin-1").strip()

    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Error parsing document content ({file.filename}): {str(e)}"
        )

    if not extracted_text or len(extracted_text.strip()) < 5:
        raise HTTPException(
            status_code=400,
            detail=f"Could not extract readable text from document '{file.filename}'. Please ensure the document is not an image-only scan or password protected."
        )

    # 4. Pass through AI extraction workflow
    try:
        return process_complaint_text_with_ai(extracted_text)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"AI extraction failed for uploaded file: {str(e)}"
        )


@app.post("/complaint/edit")
def edit_complaint(request: EditComplaintRequest):

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "system",
                "content": """
You are an AI assistant for an enterprise pharmaceutical quality assurance customer complaint management system.

You are given the current complaint data and a user's correction or additional instructions.

Update ONLY the fields affected by the correction.
Keep all other existing fields unchanged.

After updating the complaint, reassess the risk severity, priority, and suggested next action.

Return ONLY valid JSON.
Do not use Markdown.
Do not add explanations outside the JSON.

Return exactly this structure:

{
    "complaint_source": "",
    "customer_name": "",
    "product_name": "",
    "strength": "",
    "batch_number": "",
    "manufacturing_date": "",
    "expiry_date": "",
    "affected_quantity": "",
    "complaint_type": "",
    "complaint_date": "",
    "complaint_description": "",
    "risk_assessment": {
        "severity": "",
        "priority": "",
        "initial_risk": "",
        "suggested_next_action": ""
    }
}

Severity must be one of: Low, Medium, High, Critical.
Priority must be one of: P1 - Critical, P2 - High, P3 - Medium, P4 - Low.

Do not invent information.
"""
            },
            {
                "role": "user",
                "content": f"""
Current complaint data:

{request.current_data}

User correction / query:

{request.correction}
"""
            }
        ],
        response_format={"type": "json_object"}
    )

    result = response.choices[0].message.content.strip()
    if result.startswith("```json"):
        result = result[7:]
    if result.startswith("```"):
        result = result[3:]
    if result.endswith("```"):
        result = result[:-3]

    return json.loads(result.strip())